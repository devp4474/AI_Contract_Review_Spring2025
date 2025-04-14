from transformers import AutoTokenizer, AutoModelForMaskedLM, pipeline
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import fitz  # PyMuPDF
import re
import os
import nltk
from nltk.tokenize import PunktSentenceTokenizer
import pickle

# 🔐 Force NLTK to use only your punkt directory
nltk_data_path = os.path.expanduser('~/nltk_data')
nltk.data.path.clear()
nltk.data.path.append(nltk_data_path)

# Load tokenizer directly from english.pickle (avoids punkt_tab errors)
# Load English tokenizer directly from pickle (avoids all internal NLTK loading logic)
punkt_path = os.path.join(nltk_data_path, "tokenizers", "punkt", "english.pickle")
with open(punkt_path, "rb") as f:
    tokenizer = pickle.load(f)

def sent_tokenize(text):
    return tokenizer.tokenize(text)


# Clean and XML-safe text
def clean_text(text):
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]", "", text)
    text = re.sub(r"[\uD800-\uDFFF]", "", text)
    text = text.replace("\uFFFE", "").replace("\uFEFF", "")
    return text.strip()

# Load LegalBERT
tokenizer_bert = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased")
model = AutoModelForMaskedLM.from_pretrained("nlpaueb/legal-bert-base-uncased")
fill_mask = pipeline("fill-mask", model=model, tokenizer=tokenizer_bert)

# Keyword triggers
PROBLEM_KEYWORDS = [
    "indemnification", "liability", "termination", "penalty",
    "warranty", "damages", "breach", "limit", "governing law", "default"
]

def extract_text_from_pdf_or_docx(file_path):
    if file_path.endswith(".pdf"):
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return clean_text(text)
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return clean_text("\n".join([para.text for para in doc.paragraphs]))
    return ""

def analyze_contract_with_legalbert(contract_text):
    paragraphs = [p.strip() for p in contract_text.split("\n\n") if len(p.strip()) > 40]
    flagged_clauses = []

    for i, para in enumerate(paragraphs):
        clean_para = clean_text(re.sub(r"\s+", " ", para))
        if any(keyword in clean_para.lower() for keyword in PROBLEM_KEYWORDS):
            suggestion = generate_alternative_with_legalbert(clean_para)
            flagged_clauses.append({
                "clause_id": f"Clause {i + 1}",
                "original_text": clean_para,
                "issue": "Potential risky clause based on keyword match.",
                "suggestion": suggestion
            })

    return {"flagged_clauses": flagged_clauses}

def generate_alternative_with_legalbert(text):
    masked = text.replace("shall", "[MASK]", 1) if "shall" in text else f"[MASK] {text}"
    try:
        suggestions = fill_mask(masked)
        return suggestions[0]['sequence']
    except:
        return "Consider rephrasing this clause for clarity or reduced risk."

def add_comment(run, text):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    run.add_break()
    comment = run.insert_paragraph_after(f"📝 AI Note: {text}")
    comment.runs[0].font.size = Pt(9)
    comment.runs[0].font.italic = True
    comment.runs[0].font.color.rgb = RGBColor(0, 102, 204)

def create_annotated_contract(original_text, ai_output, output_path):
    doc = Document()
    doc.add_heading("AI Annotated Contract Review", level=1)

    flagged = ai_output.get("flagged_clauses", [])
    flagged_texts = [clean_text(c["original_text"]) for c in flagged]

    sentences = sent_tokenize(original_text)

    for sentence in sentences:
        clean_sentence = clean_text(sentence)

        match = next((c for c in flagged if clean_text(c["original_text"]) in clean_sentence), None)
        if match:
            p = doc.add_paragraph()
            run = p.add_run(clean_sentence)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True

            note = doc.add_paragraph()
            note_run = note.add_run(f"🧠 AI Note:\nRisk: {match['issue']}\nSuggestion: {match['suggestion']}")
            note_run.font.color.rgb = RGBColor(0, 102, 204)
            note_run.font.italic = True
        else:
            doc.add_paragraph(clean_sentence)

    doc.save(output_path)
